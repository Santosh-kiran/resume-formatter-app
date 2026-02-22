from fastapi import FastAPI, UploadFile, File
from fastapi.responses import FileResponse
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import shutil
import os
import re
import tempfile

app = FastAPI()

def format_name(name):
    return name.strip().title()

def clean_text(text):
    text = re.sub(r'[•●▪■]', '', text)
    text = re.sub(r'http\S+', '', text)
    return text.strip()

def add_zero_spacing(paragraph):
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

@app.post("/upload")
async def upload_resume(file: UploadFile = File(...)):

    temp_dir = tempfile.mkdtemp()
    input_path = os.path.join(temp_dir, file.filename)

    with open(input_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    # For demo — assume plain text extraction
    with open(input_path, "r", errors="ignore") as f:
        content = f.read()

    # Extract candidate name (first line)
    lines = content.split("\n")
    candidate_name = format_name(lines[0])

    doc = Document()

    # Name formatting
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_para.add_run(candidate_name)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    add_zero_spacing(name_para)

    doc.add_paragraph("")

    # Summary Section
    summary_heading = doc.add_paragraph("Summary")
    summary_heading.runs[0].bold = True
    summary_heading.runs[0].font.name = "Times New Roman"
    summary_heading.runs[0].font.size = Pt(10)
    add_zero_spacing(summary_heading)

    for line in lines[1:5]:
        clean_line = clean_text(line)
        para = doc.add_paragraph(f"• {clean_line}")
        para.runs[0].font.name = "Times New Roman"
        para.runs[0].font.size = Pt(10)
        add_zero_spacing(para)

    output_filename = f"{candidate_name}.docx"
    output_path = os.path.join(temp_dir, output_filename)

    doc.save(output_path)

    return FileResponse(
        output_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=output_filename
    )